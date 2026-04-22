import io
import os
import re

from docx import Document as DocxDocument
from docx.document import Document as _DocxDocumentType
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph

from app.parsers.base import CandidateItem, DocumentParser, ParsedDocument, build_fragment_ref

DOCX_CONTEXT_DEBUG = os.getenv("SAP_DOCX_CONTEXT_DEBUG", "").lower() in {"1", "true", "yes", "on"}
_SECTION_LABEL_PATTERN = re.compile(
    r"\b("
    r"sap\s+mm|materials?\s+management|"
    r"sap\s+sd|sales\s+and\s+distribution|"
    r"sap\s+fi|financial\s+accounting|"
    r"sap\s+co|controlling|"
    r"sap\s+pp|production\s+planning|"
    r"sap\s+pm|plant\s+maintenance|"
    r"sap\s+ps|project\s+system|"
    r"sap\s+wm|warehouse\s+management|"
    r"sap\s+hcm|human\s+resources|human\s+capital\s+management|hr"
    r")\b",
    re.IGNORECASE,
)


def _iter_block_items(parent: _DocxDocumentType | _Cell):
    parent_elm = parent.element.body if isinstance(parent, _DocxDocumentType) else parent._tc
    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)


def _heading_level(paragraph: Paragraph) -> int | None:
    style_name = (paragraph.style.name or "").strip().lower()
    if not style_name.startswith("heading"):
        return None
    parts = style_name.split()
    if not parts:
        return None
    maybe_level = parts[-1]
    if maybe_level.isdigit():
        return int(maybe_level)
    return 1


def _clean_cell(value: str) -> str:
    return " ".join(value.split()).strip()


def _looks_like_section_heading(text: str) -> bool:
    cleaned = _clean_cell(text)
    if not cleaned:
        return False
    if len(cleaned) > 180:
        return False
    if "|" in cleaned:
        return False
    return _SECTION_LABEL_PATTERN.search(cleaned) is not None


class DocxParser(DocumentParser):
    def parse(self, source: str | bytes) -> ParsedDocument:
        doc = DocxDocument(io.BytesIO(source)) if isinstance(source, bytes) else DocxDocument(source)
        candidates: list[CandidateItem] = []
        raw_text_chunks: list[str] = []
        fragment_idx = 1
        context_blocks: list[dict] = []
        heading_stack: list[tuple[int, str]] = []
        last_section_title: str | None = None
        last_parent_titles: list[str] = []

        for block in _iter_block_items(doc):
            if isinstance(block, Paragraph):
                text = block.text.strip()
                if not text:
                    continue
                level = _heading_level(block)
                inferred_heading = False
                if level is None and _looks_like_section_heading(text):
                    # Some business documents use normal paragraphs for section titles.
                    # Treat explicit module section labels as headings for context propagation.
                    level = 2
                    inferred_heading = True
                if level is not None:
                    heading_stack = [(lvl, title) for (lvl, title) in heading_stack if lvl < level]
                    heading_stack.append((level, text))
                section_title = heading_stack[-1][1] if heading_stack else last_section_title
                parent_titles = [title for _, title in heading_stack[:-1]] if len(heading_stack) > 1 else last_parent_titles
                if section_title:
                    last_section_title = section_title
                    last_parent_titles = list(parent_titles)
                context_blocks.append(
                    {
                        "kind": "text",
                        "text": text,
                        "section_title": section_title,
                        "parent_section_titles": parent_titles,
                        "is_heading": level is not None,
                        "is_inferred_heading": inferred_heading,
                    }
                )
                continue

            if isinstance(block, Table):
                rows: list[list[str]] = []
                for row in block.rows:
                    values = [_clean_cell(cell.text) for cell in row.cells]
                    if any(values):
                        rows.append(values)
                if not rows:
                    continue
                header_row = rows[0]
                has_header = any(header_row)
                section_title = heading_stack[-1][1] if heading_stack else last_section_title
                parent_titles = [title for _, title in heading_stack[:-1]] if len(heading_stack) > 1 else last_parent_titles
                if not section_title:
                    # Fallback for continued/fragmented table blocks where heading styling was lost.
                    header_text = " ".join(value for value in header_row if value).strip()
                    if _looks_like_section_heading(header_text):
                        section_title = header_text
                        parent_titles = []
                        last_section_title = section_title
                        last_parent_titles = []
                if section_title:
                    last_section_title = section_title
                    last_parent_titles = list(parent_titles)
                for row_idx, row_values in enumerate(rows):
                    text_values = [value for value in row_values if value]
                    if not text_values:
                        continue
                    row_text = " | ".join(text_values)
                    row_map: dict[str, str] = {}
                    if has_header and row_idx > 0:
                        limit = min(len(header_row), len(row_values))
                        for idx in range(limit):
                            header = header_row[idx].strip()
                            value = row_values[idx].strip()
                            if header and value:
                                row_map[header] = value
                    context_blocks.append(
                        {
                            "kind": "table",
                            "text": row_text,
                            "section_title": section_title,
                            "parent_section_titles": parent_titles,
                            "table_context": {
                                "is_structured_table": True,
                                "column_headers": [header for header in header_row if header],
                                "row_values": row_map,
                                "row_index": row_idx,
                            },
                        }
                    )
                    if DOCX_CONTEXT_DEBUG and "AFPO" in row_text.upper():
                        print(
                            "DEBUG: DOCX AFPO row parsed "
                            f"row='{row_text}' section='{section_title}' parents={parent_titles}",
                            flush=True,
                        )

        for idx, block in enumerate(context_blocks):
            text = block["text"]
            before = context_blocks[idx - 1]["text"] if idx > 0 else ""
            after = context_blocks[idx + 1]["text"] if idx + 1 < len(context_blocks) else ""
            nearby_text = " ".join(part for part in (before, after) if part).strip()
            extracted_data = {
                "fragment_kind": "table" if block["kind"] == "table" else "text",
                "section_title": block.get("section_title"),
                "parent_section_titles": block.get("parent_section_titles", []),
                "nearby_text": nearby_text,
            }
            if block["kind"] == "text":
                extracted_data["is_heading"] = bool(block.get("is_heading"))
                extracted_data["is_inferred_heading"] = bool(block.get("is_inferred_heading"))
            if block["kind"] == "table":
                extracted_data["table_context"] = block.get("table_context", {})
                if DOCX_CONTEXT_DEBUG and "AFPO" in text.upper():
                    print(
                        "DEBUG: DOCX AFPO extracted_data "
                        f"section_title='{extracted_data.get('section_title')}'",
                        flush=True,
                    )

            raw_text_chunks.append(text)
            candidates.append(
                CandidateItem(
                    item_type="table_row" if block["kind"] == "table" else "text_snippet",
                    title="Table fragment" if block["kind"] == "table" else "Text fragment",
                    content=text,
                    source_ref=build_fragment_ref(fragment_idx),
                    confidence=None,
                    extracted_data=extracted_data,
                )
            )
            fragment_idx += 1

        return ParsedDocument(raw_text="\n".join(raw_text_chunks), candidates=candidates)
